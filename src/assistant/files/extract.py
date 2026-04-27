"""Phase 6a — file-attachment extractors.

Each public extractor takes a :class:`pathlib.Path` to a downloaded
attachment, returns ``(extracted_text, char_count)``, or raises
:class:`ExtractionError`.

Caps:

* DOCX — no per-doc cap; the post-extract 200K char cap in the handler
  is sufficient. Document-order traversal (``<w:p>`` / ``<w:tbl>``
  iteration) preserves the source reading order — naive
  ``doc.paragraphs`` then ``doc.tables`` mangles documents that
  interleave paragraph/table/paragraph blocks (RQ4 spike).
* XLSX — ``ROW_CAP=20``, ``COL_CAP=30`` per sheet (Q13 retune; ~3 s
  wall-clock at 20 MB pre-download cap, 42 MB peak RSS at 308 MB
  worst-case input). Total-cap ``POST_EXTRACT_CHAR_CAP=200_000``
  enforced on a clean sheet boundary; defensive final substring assert
  lives in the handler.
* TXT/MD — ``utf-8-sig`` decode auto-strips a leading BOM.
* PDF — fallback path. Used when SDK ``Read`` tool's multimodal PDF
  payload doesn't propagate over the OAuth-CLI auth path. Sub-100-char
  total → Russian "no text layer" hint.

Error policy: every extractor catches a broad ``Exception`` at the leaf
parser call and re-raises as :class:`ExtractionError`. Narrow catches
(``except PackageNotFoundError`` only) leak ``KeyError`` /
``ValueError`` — both observed in the wild for mangled XML namespaces
and custom XML — and the handler can only discriminate on
``ExtractionError``.
"""

from __future__ import annotations

import contextlib
from collections.abc import Callable
from pathlib import Path

from assistant.logger import get_logger

log = get_logger("files.extract")

# Public caps. Coder/reviewer tests pin these explicitly; bumping any of
# the three requires a matching docs + plan update.
POST_EXTRACT_CHAR_CAP = 200_000
XLSX_ROW_CAP = 20
XLSX_COL_CAP = 30
PDF_MIN_TEXT_LAYER_CHARS = 100

# Russian hint surfaced by extract_pdf when the text layer is missing.
# Stable string — handler does not branch on it; the model sees it as
# the user-facing extracted-text payload.
PDF_NO_TEXT_LAYER_HINT = (
    "[PDF appears to have no text layer; OCR not available in phase 6a]"
)


class ExtractionError(Exception):
    """Raised by an extractor when the file is unreadable.

    Reason categories used by the handler for user-facing messages:

    * ``"encrypted"`` — password-protected file (DOCX/PDF). Handler
      replies in Russian: *"файл зашифрован — пришли расшифрованный"*.
    * ``"corrupt …"`` — truncated/malformed file. Handler replies
      *"не смог прочитать файл: …"*.

    The handler runs ``"encrypted" in str(exc).lower()`` to discriminate
    — keep that substring in the message text.
    """


# ----------------------------------------------------------------------
# DOCX — document-order traversal
# ----------------------------------------------------------------------


def extract_docx(path: Path) -> tuple[str, int]:
    """Extract DOCX in source-document order.

    Walks ``<w:p>`` (paragraphs) and ``<w:tbl>`` (tables) children of
    ``doc.element.body`` in iteration order. Naive ``doc.paragraphs``-
    then-``doc.tables`` extraction emits all paragraph text first then
    all table content — the model sees re-ordered narrative. RQ4 spike
    confirmed 100 % char recall on a synthetic Cyrillic + headings +
    bullet/numbered + 4x4 table + bold/italic/superscript fixture.

    Caveats not handled (flag if owner reports issues): tracked changes
    (``<w:ins>``/``<w:del>``), real footnotes (separate XML part),
    comments, hyperlinks, embedded images, SmartArt, equations.
    """
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
        from docx.oxml.ns import qn
    except ImportError as exc:  # pragma: no cover - dep-presence guard
        raise ExtractionError(f"python-docx not installed: {exc}") from exc

    try:
        doc = Document(str(path))
    except PackageNotFoundError as exc:
        # python-docx raises PackageNotFoundError both for "not a zip"
        # AND for password-protected DOCX (docx encryption rewraps the
        # OPC package as an OLE compound binary that the zipfile-based
        # loader can't parse). We surface a single message with the
        # "encrypted" keyword so the handler's Russian reply path fires.
        raise ExtractionError("encrypted or corrupted DOCX") from exc
    except (KeyError, ValueError) as exc:
        # Mangled XML namespaces / custom XML / unexpected schema
        # (devil M7).
        raise ExtractionError(f"corrupt DOCX: {exc}") from exc
    except Exception as exc:  # belt-and-suspenders for unknown errors
        raise ExtractionError(f"DOCX parse failed: {exc}") from exc

    out: list[str] = []
    try:
        body = doc.element.body
        p_tag = qn("w:p")
        tbl_tag = qn("w:tbl")
        t_tag = qn("w:t")
        tr_tag = qn("w:tr")
        tc_tag = qn("w:tc")

        for child in body.iterchildren():
            if child.tag == p_tag:
                # Concatenate all w:t runs to recover paragraph text;
                # python-docx silently drops style runs which is what
                # we want.
                text = "".join(t.text or "" for t in child.iter(t_tag))
                if text.strip():
                    out.append(text)
            elif child.tag == tbl_tag:
                for row in child.iter(tr_tag):
                    cells = [
                        "".join(t.text or "" for t in cell.iter(t_tag))
                        for cell in row.iter(tc_tag)
                    ]
                    out.append("\t".join(cells))
            # Other body children (sectPr, sdt, etc.) are intentionally
            # ignored — they carry no user-readable text.
    except Exception as exc:
        raise ExtractionError(f"DOCX traversal failed: {exc}") from exc

    text = "\n".join(out)
    return text, len(text)


# ----------------------------------------------------------------------
# XLSX — caps + clean-boundary truncation
# ----------------------------------------------------------------------


def extract_xlsx(path: Path) -> tuple[str, int]:
    """Extract XLSX with row/col caps and total-cap clean-boundary.

    All sheets walked; each sheet capped at ``XLSX_ROW_CAP`` rows by
    ``XLSX_COL_CAP`` cols. Cells joined with ``\\t``; rows joined with
    ``\\n``; sheets separated by ``\\n\\n`` and prefixed
    ``=== Sheet: <name> ===``.

    Total cap (``POST_EXTRACT_CHAR_CAP``) enforced on a clean sheet
    boundary inside this function — the handler runs a defensive final
    substring truncation to handle the worst case where a single sheet
    exceeds the cap on its own (~6700 capped rows by 30 cells by 1 char,
    well above ``XLSX_ROW_CAP``).

    Mandatory openpyxl flags (RQ2):

    * ``read_only=True`` — keeps peak RSS ≈ 42 MB on a 308 MB workbook.
    * ``data_only=True`` — reads cached formula results; raw XLSX
      generated without an Excel save pass shows ``None`` for formula
      cells (devil L3).
    * ``wb.close()`` in ``finally`` — openpyxl leaks fd otherwise.
    """
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover - dep-presence guard
        raise ExtractionError(f"openpyxl not installed: {exc}") from exc

    # zipfile is in stdlib; import is cheap and unconditional.
    from zipfile import BadZipFile

    try:
        wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    except BadZipFile as exc:
        raise ExtractionError("encrypted or corrupted XLSX") from exc
    except KeyError as exc:
        # openpyxl raises KeyError for missing required parts (corrupt
        # zip without the workbook.xml entry).
        raise ExtractionError(f"corrupt XLSX: {exc}") from exc
    except Exception as exc:
        raise ExtractionError(f"XLSX parse failed: {exc}") from exc

    out: list[str] = []
    total_chars = 0
    sheet_names = list(wb.sheetnames)
    skipped_sheets = 0
    try:
        for sheet_idx, sheet_name in enumerate(sheet_names):
            chunk_lines: list[str] = [f"=== Sheet: {sheet_name} ==="]
            ws = wb[sheet_name]
            # Defense-in-depth row cap: ``max_row=XLSX_ROW_CAP`` is the
            # primary control, but earlier openpyxl 3.1.x sub-versions
            # ignore ``max_row`` on ``read_only=True`` workbooks past
            # the first row (review code M3). The pyproject pin
            # (``openpyxl>=3.1.5``) avoids the bug, and the explicit
            # ``enumerate`` + early ``break`` here ensures a future
            # regression in the dependency cannot blow the RQ2 RSS
            # guarantee.
            for row_idx, row in enumerate(
                ws.iter_rows(
                    values_only=True,
                    max_row=XLSX_ROW_CAP,
                    max_col=XLSX_COL_CAP,
                )
            ):
                if row_idx >= XLSX_ROW_CAP:
                    break
                cells = [
                    "" if v is None else str(v) for v in row[:XLSX_COL_CAP]
                ]
                chunk_lines.append("\t".join(cells))
            chunk = "\n".join(chunk_lines)
            # Sheet-boundary total cap (devil H1).
            # ``+ 2`` accounts for the ``\n\n`` separator between
            # sheets in the final ``"\n\n".join(out)``.
            tentative = total_chars + len(chunk) + (2 if out else 0)
            if tentative > POST_EXTRACT_CHAR_CAP:
                skipped_sheets = len(sheet_names) - sheet_idx
                break
            out.append(chunk)
            total_chars = tentative

        if skipped_sheets > 0:
            out.append(
                f"[…truncated; {skipped_sheets} more sheet(s) skipped at "
                f"{POST_EXTRACT_CHAR_CAP}-char cap]"
            )
    except Exception as exc:
        raise ExtractionError(f"XLSX traversal failed: {exc}") from exc
    finally:
        # openpyxl read-only workbooks leak fd if not explicitly closed.
        with contextlib.suppress(Exception):  # pragma: no cover
            wb.close()

    text = "\n\n".join(out)
    return text, len(text)


# ----------------------------------------------------------------------
# TXT / MD
# ----------------------------------------------------------------------


def extract_txt(path: Path) -> tuple[str, int]:
    """Extract a UTF-8 text file.

    ``encoding="utf-8-sig"`` auto-strips a leading UTF-8 BOM (``\\ufeff``).
    Plain ``utf-8`` would leave the BOM as a literal character at the
    head of the string (devil L2).

    ``errors="replace"`` substitutes ``\\ufffd`` for any invalid byte
    sequence — the model handles the replacement marker fine and the
    extractor never raises on encoding noise alone.
    """
    try:
        text = path.read_text(encoding="utf-8-sig", errors="replace")
    except OSError as exc:
        raise ExtractionError(f"read failed: {exc}") from exc
    return text, len(text)


# ``extract_md`` is byte-for-byte identical to ``extract_txt``: the
# handler's downstream context-prep treats the markdown content as
# plain text.
extract_md = extract_txt


# ----------------------------------------------------------------------
# PDF — fallback / pypdf-uniform path
# ----------------------------------------------------------------------


def extract_pdf(path: Path) -> tuple[str, int]:
    """Extract a PDF via ``pypdf`` (fallback / pypdf-uniform path).

    Used when the SDK Read tool's multimodal payload does NOT propagate
    over the OAuth-CLI auth path (RQ1 live probe FAIL case). The
    ``EXTRACTORS`` dispatch table includes ``"pdf"`` so the handler can
    transparently fall back by flipping a single discriminator.

    Sub-``PDF_MIN_TEXT_LAYER_CHARS`` total → return the
    ``PDF_NO_TEXT_LAYER_HINT`` string with ``char_count=0``. Devil M8
    accepted: a real text PDF with <100 chars total triggers the same
    hint (single-user trust model + image-PDF being the dominant case
    for short-text outputs).
    """
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError
    except ImportError as exc:  # pragma: no cover - dep-presence guard
        raise ExtractionError(f"pypdf not installed: {exc}") from exc

    try:
        reader = PdfReader(str(path))
    except PdfReadError as exc:
        raise ExtractionError(f"corrupt PDF: {exc}") from exc
    except Exception as exc:
        raise ExtractionError(f"PDF parse failed: {exc}") from exc

    if reader.is_encrypted:
        raise ExtractionError("encrypted PDF")

    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            # One bad page should not abort the whole extract; the
            # model can usually answer from the remaining pages.
            t = ""
        parts.append(t)
    text = "\n\n".join(parts).strip()
    if len(text) < PDF_MIN_TEXT_LAYER_CHARS:
        return (PDF_NO_TEXT_LAYER_HINT, 0)
    return text, len(text)


# ----------------------------------------------------------------------
# Public dispatch table
# ----------------------------------------------------------------------

# Handler keys on the attachment kind (``IncomingMessage.attachment_kind``)
# to pick the extractor. The pdf entry is only consumed when the PDF
# native-Read path is disabled (Option C live probe FAIL or Mac dev
# fallback).
EXTRACTORS: dict[str, Callable[[Path], tuple[str, int]]] = {
    "docx": extract_docx,
    "xlsx": extract_xlsx,
    "txt": extract_txt,
    "md": extract_md,
    "pdf": extract_pdf,
}
