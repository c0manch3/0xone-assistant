"""RQ4 — python-docx Cyrillic + complex layout fidelity.

No real DOCX corpus on this machine, so we **synthesize** a complex
DOCX with python-docx (Russian text, mixed bold/italic, headings,
bullet/numbered lists, a 3x4 table with Cyrillic cells, and a
footnote-like superscripted run), keep the **exact source text** as
ground truth, then re-extract via python-docx and measure
character-level recall.

Caveat: this is a synthesis-and-roundtrip test. It bounds python-docx's
capability for the layout features we generate. Real owner DOCX files
may include features we don't synthesize (revisions, comments,
SmartArt, real footnotes, equations) — flagged in findings.

PASS: ≥ 95 % character recall.
FAIL: plumb pandoc.

Run:
    /tmp/.spike6a-venv/bin/python plan/phase6a/spikes/rq4_docx_cyrillic.py
"""

from __future__ import annotations

import sys
import tempfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


# ---------------------------------------------------------------------------
# Source content — single source of truth for the ground-truth string.
# We store it as a list of (kind, payload) records so generation and
# ground-truth derivation stay in lockstep.
# ---------------------------------------------------------------------------
SOURCE: list[tuple[str, object]] = [
    ("heading", "Отчёт о квартальных показателях"),
    ("para", "Это первый абзац документа на русском языке. Здесь есть "
             "обычный текст, жирный и курсив."),
    ("para_runs", [
        ("Смешанные стили: ", None),
        ("жирный фрагмент", "bold"),
        (", ", None),
        ("курсив", "italic"),
        (", и обычный текст. Сноска", None),
        ("¹", "superscript"),
        (" в конце.", None),
    ]),
    ("heading", "Список достижений"),
    ("bullet", "Запустили новую систему мониторинга."),
    ("bullet", "Сократили время отклика на 30%."),
    ("bullet", "Добавили поддержку Юникода и эмодзи."),
    ("heading", "Шаги внедрения"),
    ("number", "Анализ требований и сбор обратной связи."),
    ("number", "Разработка прототипа."),
    ("number", "Тестирование и приёмка."),
    ("heading", "Сводная таблица"),
    ("table", [
        ["Квартал", "Выручка", "Прибыль", "Комментарий"],
        ["Q1", "1 200 000", "180 000", "Стабильный рост"],
        ["Q2", "1 450 000", "210 000", "Запуск новинок"],
        ["Q3", "1 380 000", "195 000", "Сезонный спад"],
    ]),
    ("para", "Заключительный параграф со специальными символами: "
             "«ёлки-палки», тире — и многоточие… Конец."),
]


def build_docx(path: Path) -> str:
    """Generate the test DOCX. Returns the ground-truth canonical string."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    gt_lines: list[str] = []

    for kind, payload in SOURCE:
        if kind == "heading":
            assert isinstance(payload, str)
            doc.add_heading(payload, level=1)
            gt_lines.append(payload)
        elif kind == "para":
            assert isinstance(payload, str)
            p = doc.add_paragraph(payload)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            gt_lines.append(payload)
        elif kind == "para_runs":
            assert isinstance(payload, list)
            p = doc.add_paragraph()
            line_parts: list[str] = []
            for text, fmt in payload:
                run = p.add_run(text)
                if fmt == "bold":
                    run.bold = True
                elif fmt == "italic":
                    run.italic = True
                elif fmt == "superscript":
                    run.font.superscript = True
                line_parts.append(text)
            gt_lines.append("".join(line_parts))
        elif kind == "bullet":
            assert isinstance(payload, str)
            doc.add_paragraph(payload, style="List Bullet")
            gt_lines.append(payload)
        elif kind == "number":
            assert isinstance(payload, str)
            doc.add_paragraph(payload, style="List Number")
            gt_lines.append(payload)
        elif kind == "table":
            assert isinstance(payload, list)
            rows = payload
            tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
            for r_idx, row in enumerate(rows):
                cells = tbl.rows[r_idx].cells
                for c_idx, cell_text in enumerate(row):
                    cells[c_idx].text = cell_text
                gt_lines.append("\t".join(row))
        else:
            raise AssertionError(f"unknown kind {kind!r}")

    doc.save(str(path))
    return "\n".join(gt_lines)


def extract(path: Path) -> str:
    """Mirror the planned extractor: paragraphs + flattened tables."""
    doc = Document(str(path))
    out: list[str] = []
    for para in doc.paragraphs:
        text = para.text
        if text:  # drop empty paragraphs (matches plan's "join non-empty")
            out.append(text)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_cells = [cell.text for cell in row.cells]
            out.append("\t".join(row_cells))
    return "\n".join(out)


def char_recall(gt: str, ext: str) -> float:
    """Multiset character recall: |gt ∩ ext| / |gt|.

    Uses Counter intersection so reordering, duplication of newlines, and
    extraction of MORE than gt do not hurt the score. This is a loose
    metric; we additionally check token presence.
    """
    gc, ec = Counter(gt), Counter(ext)
    inter = sum((gc & ec).values())
    return inter / max(1, sum(gc.values()))


def token_presence(gt: str, ext: str) -> tuple[int, int, list[str]]:
    """Whitespace-split tokens; report missing tokens."""
    gt_tokens = [t for t in gt.split() if t]
    missing = [t for t in gt_tokens if t not in ext]
    return len(gt_tokens) - len(missing), len(gt_tokens), missing


def main() -> int:
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "sample.docx"
        gt = build_docx(path)
        ext = extract(path)

        size = path.stat().st_size

        recall = char_recall(gt, ext)
        present, total, missing = token_presence(gt, ext)

        # Side-by-side preview (truncated).
        print("=== GROUND TRUTH (first 500 chars) ===")
        print(gt[:500])
        print()
        print("=== EXTRACTED (first 500 chars) ===")
        print(ext[:500])
        print()
        print("=== METRICS ===")
        print(f"docx size:           {size} bytes")
        print(f"gt char count:       {len(gt)}")
        print(f"ext char count:      {len(ext)}")
        print(f"char recall:         {recall * 100:.2f}%")
        print(f"token recall:        {present}/{total} ({present / total * 100:.2f}%)")
        if missing:
            print(f"missing tokens:      {missing[:20]}")

        passed = recall >= 0.95
        verdict = "PASS" if passed else "FAIL"
        print(f"VERDICT: {verdict} (target ≥ 95% char recall)")
        return 0 if passed else 1


if __name__ == "__main__":
    sys.exit(main())
