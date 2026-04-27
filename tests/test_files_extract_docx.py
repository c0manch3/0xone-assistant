"""Phase 6a — DOCX extractor unit tests.

Covers:
- happy path: synthetic DOCX with mixed paragraphs/tables → document order;
- Cyrillic + special characters preserved;
- corrupt zip → ExtractionError;
- empty document → empty string.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from assistant.files.extract import ExtractionError, extract_docx


def _make_docx(path: Path, builder) -> None:  # type: ignore[no-untyped-def]
    """Helper: build a DOCX file via python-docx."""
    from docx import Document

    doc = Document()
    builder(doc)
    doc.save(str(path))


def test_extract_docx_paragraphs_in_order(tmp_path: Path) -> None:
    """Plain paragraphs flatten to ``\\n``-joined text."""

    def build(doc):  # type: ignore[no-untyped-def]
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        doc.add_paragraph("Third paragraph.")

    p = tmp_path / "simple.docx"
    _make_docx(p, build)

    text, n = extract_docx(p)
    assert n == len(text)
    lines = text.split("\n")
    assert lines == ["First paragraph.", "Second paragraph.", "Third paragraph."]


def test_extract_docx_table_cells_tab_joined(tmp_path: Path) -> None:
    """Each table row becomes one ``\\t``-joined line."""

    def build(doc):  # type: ignore[no-untyped-def]
        table = doc.add_table(rows=2, cols=3)
        for r in range(2):
            for c in range(3):
                table.cell(r, c).text = f"r{r}c{c}"

    p = tmp_path / "table.docx"
    _make_docx(p, build)

    text, _ = extract_docx(p)
    assert "r0c0\tr0c1\tr0c2" in text
    assert "r1c0\tr1c1\tr1c2" in text


def test_extract_docx_document_order_paragraph_table_paragraph(
    tmp_path: Path,
) -> None:
    """RQ4 fix: tables emit at their source position, NOT after all paragraphs.

    With the naive paragraphs-then-tables approach, ``After`` would
    appear before the table content. Document-order traversal puts the
    table line between ``Before`` and ``After``.
    """

    def build(doc):  # type: ignore[no-untyped-def]
        doc.add_paragraph("Before")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "X"
        table.cell(0, 1).text = "Y"
        doc.add_paragraph("After")

    p = tmp_path / "order.docx"
    _make_docx(p, build)

    text, _ = extract_docx(p)
    lines = text.split("\n")
    assert lines.index("Before") < lines.index("X\tY")
    assert lines.index("X\tY") < lines.index("After")


def test_extract_docx_cyrillic_round_trip(tmp_path: Path) -> None:
    """Russian + special-quotes + em-dash survive extraction unchanged."""

    payload = "Привет, мир — это «тест» с кавычками."

    def build(doc):  # type: ignore[no-untyped-def]
        doc.add_paragraph(payload)

    p = tmp_path / "cyr.docx"
    _make_docx(p, build)

    text, _ = extract_docx(p)
    assert payload in text


def test_extract_docx_skips_empty_paragraphs(tmp_path: Path) -> None:
    """Whitespace-only paragraphs are dropped to keep envelope tight."""

    def build(doc):  # type: ignore[no-untyped-def]
        doc.add_paragraph("real text")
        doc.add_paragraph("")
        doc.add_paragraph("   ")
        doc.add_paragraph("more text")

    p = tmp_path / "blanks.docx"
    _make_docx(p, build)

    text, _ = extract_docx(p)
    assert text == "real text\nmore text"


def test_extract_docx_corrupt_raises(tmp_path: Path) -> None:
    """A non-DOCX file masquerading as ``.docx`` raises ExtractionError."""
    p = tmp_path / "not_really.docx"
    p.write_bytes(b"this is not a zip")

    with pytest.raises(ExtractionError) as excinfo:
        extract_docx(p)
    # Plan invariant: handler keys on the "encrypted" substring; unrelated
    # garbage gets the "encrypted or corrupted" message which still
    # surfaces the right user-facing reply.
    assert "encrypted" in str(excinfo.value).lower() or "corrupt" in str(
        excinfo.value
    ).lower()


def test_extract_docx_missing_file_raises(tmp_path: Path) -> None:
    """A nonexistent path raises ExtractionError (not FileNotFoundError)."""
    p = tmp_path / "ghost.docx"
    with pytest.raises(ExtractionError):
        extract_docx(p)


def test_extract_docx_empty_document_returns_empty_string(tmp_path: Path) -> None:
    """A DOCX with no paragraphs/tables returns ``("", 0)``."""

    def build(doc):  # type: ignore[no-untyped-def]
        # Empty document with no body content.
        return None

    p = tmp_path / "empty.docx"
    _make_docx(p, build)

    text, n = extract_docx(p)
    # Default DOCX has one empty paragraph that we strip; expect empty.
    assert text == ""
    assert n == 0
