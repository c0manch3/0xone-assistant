"""Phase 7 / commit 9 — tools/extract_doc/main.py CLI coverage.

Subprocess-based: mirrors the Bash-allowlist invocation path used at
runtime. Each test writes a tiny fixture to `tmp_path`, runs the CLI
with `sys.executable`, parses the JSON line on stdout (or stderr for
error cases), and asserts exit code + payload shape.

Covers the 5 dispatch paths (PDF / DOCX / XLSX / RTF / TXT) plus the
key validation / security guards: oversize rejection, encrypted PDF,
zip-bomb, `--pages` on non-PDF, `--max-chars` truncation.
"""

from __future__ import annotations

import json
import os
import subprocess
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from fpdf import FPDF
from openpyxl import Workbook

_CLI = Path(__file__).resolve().parents[1] / "tools" / "extract_doc" / "main.py"


def _run(
    *args: str,
    env_extra: dict[str, str] | None = None,
) -> subprocess.CompletedProcess[str]:
    env = dict(os.environ)
    if env_extra:
        env.update(env_extra)
    return subprocess.run(
        [sys.executable, str(_CLI), *args],
        env=env,
        capture_output=True,
        text=True,
        check=False,
        timeout=30,
    )


# ---- fixture builders ------------------------------------------------------


def _make_pdf(path: Path, body: str = "Hello, World!") -> None:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    # fpdf2 is latin-1 by default; keep bodies ASCII for portability.
    pdf.cell(0, 10, body)
    pdf.output(str(path))


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Docx paragraph one.")
    doc.add_paragraph("Docx paragraph two — Кириллица.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "CellA"
    table.rows[0].cells[1].text = "CellB"
    doc.save(str(path))


def _make_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "Header"
    ws["B1"] = "Value"
    ws["A2"] = "alpha"
    ws["B2"] = 42
    wb.save(str(path))


def _make_rtf(path: Path) -> None:
    # Minimal valid RTF — striprtf should extract "hello rtf".
    path.write_bytes(rb"{\rtf1\ansi hello rtf}")


def _make_txt(path: Path, body: str = "plain text body\nsecond line\n") -> None:
    path.write_text(body, encoding="utf-8")


# ---- happy paths (5 file types) --------------------------------------------


def test_extract_pdf_happy(tmp_path: Path) -> None:
    src = tmp_path / "doc.pdf"
    _make_pdf(src)
    r = _run(str(src))
    assert r.returncode == 0, r.stderr
    payload = json.loads(r.stdout)
    assert payload["ok"] is True
    assert payload["format"] == "pdf"
    assert payload["units"] >= 1
    assert payload["size_bytes"] > 0
    assert "Hello" in payload["text"]
    assert payload["truncated"] is False


def test_extract_docx_happy(tmp_path: Path) -> None:
    src = tmp_path / "doc.docx"
    _make_docx(src)
    r = _run(str(src))
    assert r.returncode == 0, r.stderr
    payload = json.loads(r.stdout)
    assert payload["format"] == "docx"
    assert "paragraph one" in payload["text"]
    assert "Кириллица" in payload["text"]
    # Table row rendered as tab-joined.
    assert "CellA\tCellB" in payload["text"]


def test_extract_xlsx_happy(tmp_path: Path) -> None:
    src = tmp_path / "doc.xlsx"
    _make_xlsx(src)
    r = _run(str(src))
    assert r.returncode == 0, r.stderr
    payload = json.loads(r.stdout)
    assert payload["format"] == "xlsx"
    assert "# sheet: Sheet1" in payload["text"]
    assert "Header\tValue" in payload["text"]
    assert "alpha\t42" in payload["text"]


def test_extract_rtf_happy(tmp_path: Path) -> None:
    src = tmp_path / "doc.rtf"
    _make_rtf(src)
    r = _run(str(src))
    assert r.returncode == 0, r.stderr
    payload = json.loads(r.stdout)
    assert payload["format"] == "rtf"
    assert "hello rtf" in payload["text"]


def test_extract_txt_happy(tmp_path: Path) -> None:
    src = tmp_path / "doc.txt"
    _make_txt(src)
    r = _run(str(src))
    assert r.returncode == 0, r.stderr
    payload = json.loads(r.stdout)
    assert payload["format"] == "txt"
    assert "plain text body" in payload["text"]
    assert payload["units"] >= 2  # multi-line


# ---- validation guards -----------------------------------------------------


def test_unsupported_suffix_rejected(tmp_path: Path) -> None:
    src = tmp_path / "notes.md"
    src.write_text("# heading\n", encoding="utf-8")
    r = _run(str(src))
    assert r.returncode == 3, r.stderr
    err = json.loads(r.stderr)
    assert err["ok"] is False
    assert "unsupported suffix" in err["error"]


def test_missing_path_rejected(tmp_path: Path) -> None:
    r = _run(str(tmp_path / "nope.pdf"))
    assert r.returncode == 3, r.stderr
    err = json.loads(r.stderr)
    assert err["ok"] is False
    assert "path resolve failed" in err["error"]


def test_directory_rejected(tmp_path: Path) -> None:
    target = tmp_path / "adir.pdf"
    target.mkdir()
    r = _run(str(target))
    assert r.returncode == 3, r.stderr
    err = json.loads(r.stderr)
    assert "not a regular file" in err["error"]


def test_oversize_input_rejected(tmp_path: Path) -> None:
    src = tmp_path / "big.txt"
    _make_txt(src, body="X" * 5000)
    r = _run(str(src), env_extra={"MEDIA_EXTRACT_MAX_INPUT_BYTES": "100"})
    assert r.returncode == 3, r.stderr
    err = json.loads(r.stderr)
    assert err["error"].startswith("input size")
    assert err["cap"] == 100


def test_pages_flag_rejected_for_non_pdf(tmp_path: Path) -> None:
    src = tmp_path / "doc.txt"
    _make_txt(src)
    r = _run(str(src), "--pages", "1-2")
    assert r.returncode == 3, r.stderr
    err = json.loads(r.stderr)
    assert "--pages is only supported" in err["error"]


def test_pages_parse_error(tmp_path: Path) -> None:
    src = tmp_path / "doc.pdf"
    _make_pdf(src)
    r = _run(str(src), "--pages", "abc")
    assert r.returncode == 2, r.stderr


def test_max_chars_truncation(tmp_path: Path) -> None:
    src = tmp_path / "doc.txt"
    _make_txt(src, body="A" * 500)
    r = _run(str(src), "--max-chars", "50")
    assert r.returncode == 0, r.stderr
    payload = json.loads(r.stdout)
    assert payload["truncated"] is True
    assert payload["chars"] == 50


def test_zip_bomb_guard_rejects_inflated_docx(tmp_path: Path) -> None:
    """Craft a ZIP whose declared uncompressed total blows past the 64 MB cap.

    We use 70 x 1 MB entries (70 MB declared, ~80 KB compressed via
    DEFLATE on zero bytes). This models the real attack: tiny file on
    disk, huge memory balloon on naive `extractall`. The guard runs on
    `ZipInfo.file_size` aggregation BEFORE any extractor touches the
    archive.
    """
    src = tmp_path / "bomb.docx"
    with zipfile.ZipFile(src, "w", zipfile.ZIP_DEFLATED) as zf:
        for i in range(70):
            zf.writestr(f"fat{i}.bin", b"\x00" * (1024 * 1024))
    r = _run(str(src))
    assert r.returncode == 3, r.stderr
    err = json.loads(r.stderr)
    assert "zip-bomb guard" in err["error"]
    # Total declared size should be reported in the payload.
    assert err["declared"] >= 64 * 1024 * 1024


def test_encrypted_pdf_rejected(tmp_path: Path) -> None:
    # Build a plain PDF, then re-emit it encrypted via pypdf.
    import pypdf

    plain = tmp_path / "plain.pdf"
    _make_pdf(plain)
    reader = pypdf.PdfReader(str(plain))
    writer = pypdf.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt("secret")
    enc = tmp_path / "enc.pdf"
    with open(enc, "wb") as fh:
        writer.write(fh)
    r = _run(str(enc))
    assert r.returncode == 3, r.stderr
    err = json.loads(r.stderr)
    assert "encrypted" in err["error"]


# ---- argparse --------------------------------------------------------------


def test_missing_positional_exit_2(tmp_path: Path) -> None:
    del tmp_path
    r = _run()
    assert r.returncode == 2


def test_help_exits_zero() -> None:
    r = _run("--help")
    assert r.returncode == 0
    assert "extract_doc" in r.stdout.lower() or "extract" in r.stdout.lower()


_ = pytest  # keep the import; useful for future parametrization
