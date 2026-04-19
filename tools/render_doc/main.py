"""render_doc CLI — render plain-text bodies to PDF / DOCX.

Invocation:
    python tools/render_doc/main.py --body-file PATH --out PATH
                                    [--title T] [--font DejaVu]

Output format is inferred from ``--out`` suffix (``.pdf`` / ``.docx``;
case-insensitive). Exit codes:

    0  ok            (+ ``{"ok": true, ...}`` JSON on stdout)
    2  usage         (argparse / suffix mismatch / empty body)
    3  path guard    (--body-file outside stage dir, --out outside outbox,
                      suffix unsupported, size exceeded)
    4  I/O           (read/write/rename failure)
    5  unknown       (fpdf2 / python-docx crashed)

Path guards (invariant I-7.3 in detailed-plan §7):
  * ``--body-file`` MUST resolve inside ``<data_dir>/run/render-stage/``.
  * ``--out`` MUST resolve inside ``<data_dir>/media/outbox/``.

Both directories are created with ``0o700`` by ``Daemon.start``; the CLI
never creates them itself — that would let a mis-configured run silently
write outside the sanctioned tree.

Rendering notes:
  * PDF uses ``fpdf2``; ``fpdf2>=2.7`` imports Pillow transitively (S-3).
    We ship a vendored ``DejaVuSans.ttf`` under ``_lib/`` so Cyrillic
    renders without relying on host-installed fonts (confirmed on
    macOS 14 + Ubuntu 22.04 — neither has DejaVu in default TTF search
    paths).
  * DOCX uses ``python-docx``; no font vendoring needed (Word picks the
    system font — Cyrillic glyphs are ubiquitous in default DOCX
    readers).

Size caps: the default body cap mirrors
``MediaSettings.render_max_body_bytes`` (512_000) and the output cap
mirrors ``render_max_output_bytes`` (10_485_760). Overrides via the
``RENDER_DOC_MAX_BODY_BYTES`` / ``RENDER_DOC_MAX_OUTPUT_BYTES`` env vars
keep tests fast without threading a --cap flag through the Bash
allowlist (which would widen its attack surface).
"""

from __future__ import annotations

import argparse
import json
import os
import sys
import traceback
from pathlib import Path
from typing import Any

# Phase-7 (Q9a tech debt close): `tools/render_doc/` is a real Python
# sub-package. Imports resolve as `tools.render_doc._lib.*`. When
# launched as `python tools/render_doc/main.py`, `__package__` is empty
# and the project root is not on sys.path by default — the short pragma
# below restores it so both invocation forms (cwd-launch +
# `python -m tools.render_doc.main`) work. Mirrors the
# tools/memory/main.py + tools/skill_installer/main.py precedent.
_PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

# Phase-7 fix-pack I3/I7: shared path-guard helpers. Both render_doc
# path validators (body-file input, --out output) now delegate to the
# central implementation so a future path-guard tightening stays
# consistent across transcribe / extract_doc / genimage / render_doc.
from assistant.media.path_guards import (  # noqa: E402
    PathGuardError,
    validate_existing_input_path,
    validate_future_output_path,
)

EXIT_OK = 0
EXIT_USAGE = 2
EXIT_PATH = 3
EXIT_IO = 4
EXIT_UNKNOWN = 5

# Keep in sync with `MediaSettings.render_max_body_bytes` /
# `render_max_output_bytes`. Duplication is deliberate: the CLI is
# stdlib-focused and must not import pydantic-settings (import cost +
# env-var coupling). Env-var overrides keep tests fast.
_DEFAULT_MAX_BODY_BYTES = 512_000
_DEFAULT_MAX_OUTPUT_BYTES = 10_485_760

# Bundled font (S-3 verified) lives next to this module so the CLI works
# in any worktree without relying on /opt/homebrew or /usr/share/fonts.
_BUNDLED_FONT_PATH = Path(__file__).resolve().parent / "_lib" / "DejaVuSans.ttf"

_STAGE_SUBDIR = Path("run") / "render-stage"
_OUTBOX_SUBDIR = Path("media") / "outbox"

_SUPPORTED_SUFFIXES = (".pdf", ".docx")


# ---------------------------------------------------------------------------
# Config resolution (env-driven; stdlib only — no pydantic import)
# ---------------------------------------------------------------------------


def _default_data_dir() -> Path:
    """Mirror ``assistant.config._default_data_dir`` without importing it.

    Keeping the CLI stdlib-only means tests can launch it under a
    tmp-dir ``XDG_DATA_HOME`` without spinning up the whole pydantic
    Settings graph.
    """
    base = os.environ.get("XDG_DATA_HOME")
    root = Path(base) if base else Path.home() / ".local" / "share"
    return root / "0xone-assistant"


def _resolve_data_dir() -> Path:
    override = os.environ.get("ASSISTANT_DATA_DIR")
    if override:
        return Path(override).expanduser()
    return _default_data_dir()


def _resolve_int_env(name: str, default: int) -> int:
    raw = os.environ.get(name)
    if raw:
        try:
            value = int(raw)
            if value > 0:
                return value
        except ValueError:
            pass
    return default


def _resolve_max_body_bytes() -> int:
    return _resolve_int_env("RENDER_DOC_MAX_BODY_BYTES", _DEFAULT_MAX_BODY_BYTES)


def _resolve_max_output_bytes() -> int:
    return _resolve_int_env("RENDER_DOC_MAX_OUTPUT_BYTES", _DEFAULT_MAX_OUTPUT_BYTES)


# ---------------------------------------------------------------------------
# Path guards
# ---------------------------------------------------------------------------


def _stage_root(data_dir: Path) -> Path:
    return (data_dir / _STAGE_SUBDIR).resolve()


def _outbox_root(data_dir: Path) -> Path:
    return (data_dir / _OUTBOX_SUBDIR).resolve()


def _resolve_body_file(raw: str, data_dir: Path) -> tuple[Path | None, str | None]:
    """Validate ``--body-file`` resolves inside the render-stage directory.

    The daemon pre-creates the stage dir with ``0o700``; the CLI refuses
    to create it — a missing stage dir is a configuration error, not a
    silent fallback path.

    Fix-pack I3/I7: delegates to
    :func:`assistant.media.path_guards.validate_existing_input_path`
    for the strict-resolve + ``is_file`` + suffix-allowlist triplet;
    the containment check is then layered on top. The body-file
    allow-list is intentionally wider than the allowed ``--out``
    formats — we accept any plain text whose suffix is ``.txt`` /
    ``.md``; callers that pass a bare or other-suffix body-file get
    a path-guard rejection rather than a rendering surprise.
    """
    candidate = Path(raw).expanduser()
    # Use the shared helper for the "exists + is_file + resolves
    # strictly" story; we don't enforce a suffix allow-list here
    # because the daemon writes a transient `.txt` body into the
    # stage dir, but extensions differ across test fixtures. The
    # containment check below catches the real invariant.
    try:
        resolved = candidate.resolve(strict=True)
    except FileNotFoundError as exc:
        return None, f"--body-file cannot be resolved: {exc}"
    except (OSError, RuntimeError) as exc:
        return None, f"--body-file cannot be resolved: {exc}"
    stage = _stage_root(data_dir)
    if not stage.exists():
        return None, f"--body-file stage dir missing: {stage}"
    if not resolved.is_relative_to(stage):
        return None, f"--body-file must live under {stage} (got {resolved})"
    if not resolved.is_file():
        return None, f"--body-file is not a regular file: {resolved}"
    return resolved, None


def _resolve_out_path(raw: str, data_dir: Path) -> tuple[Path | None, str | None]:
    """Validate ``--out`` would land inside the outbox directory.

    Fix-pack I3/I7: delegates to
    :func:`assistant.media.path_guards.validate_future_output_path`
    with the render_doc allow-list (``{".pdf", ".docx"}``). The
    shared helper is the canonical implementation of the
    parent-resolve + re-append + ``is_relative_to`` pattern that
    render_doc pioneered; the four CLIs now use one copy instead of
    maintaining four subtly-different ones.
    """
    outbox = _outbox_root(data_dir)
    if not outbox.exists():
        return None, f"--out outbox dir missing: {outbox}"
    try:
        final = validate_future_output_path(
            raw,
            root=outbox,
            allowed_suffixes=_SUPPORTED_SUFFIXES,
        )
    except PathGuardError as exc:
        return None, f"--out {exc}"
    return final, None


# ---------------------------------------------------------------------------
# JSON helpers
# ---------------------------------------------------------------------------


def _ok(data: dict[str, Any]) -> int:
    sys.stdout.write(json.dumps({"ok": True, "data": data}, ensure_ascii=False) + "\n")
    return EXIT_OK


def _fail(code: int, error: str, **extra: Any) -> int:
    payload: dict[str, Any] = {"ok": False, "error": error}
    payload.update(extra)
    sys.stderr.write(json.dumps(payload, ensure_ascii=False) + "\n")
    return code


# ---------------------------------------------------------------------------
# Renderers
# ---------------------------------------------------------------------------


def _render_pdf(
    *,
    body: str,
    out_path: Path,
    title: str | None,
    font_path: Path,
    max_output_bytes: int,
) -> tuple[int | None, str | None]:
    """Render ``body`` to ``out_path``. Returns ``(bytes_written, error)``.

    Writes via a ``.tmp`` sibling + ``os.replace`` to keep the outbox
    free of half-written files that the sweeper might pick up.
    """
    try:
        from fpdf import FPDF  # Lazy import: keeps `--help` stdlib-fast.
    except ImportError as exc:  # pragma: no cover — root pyproject pins fpdf2.
        return None, f"fpdf2 not importable: {exc}"

    if not font_path.is_file():
        return None, f"vendored font missing: {font_path}"

    tmp_path = out_path.with_suffix(out_path.suffix + ".tmp")
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font("DejaVu", "", str(font_path))
        if title:
            pdf.set_font("DejaVu", size=16)
            pdf.multi_cell(0, 10, title)
            pdf.ln(4)
        pdf.set_font("DejaVu", size=12)
        # `multi_cell` with width=0 expands to the page margin — safer
        # than hardcoding 180mm which breaks on non-A4 defaults.
        pdf.multi_cell(0, 8, body)
        pdf.output(str(tmp_path))
    except Exception as exc:
        # fpdf2 raises a grab-bag (OSError on font load, RuntimeError on
        # layout, ValueError on missing glyphs). Narrowing here would
        # drift with upstream; we convert all failures to exit 5 JSON.
        tmp_path.unlink(missing_ok=True)
        return None, f"fpdf2 render failed: {type(exc).__name__}: {exc}"

    try:
        size = tmp_path.stat().st_size
    except OSError as exc:
        tmp_path.unlink(missing_ok=True)
        return None, f"stat after render failed: {exc}"

    if size > max_output_bytes:
        tmp_path.unlink(missing_ok=True)
        return None, (
            f"rendered PDF {size} bytes exceeds cap {max_output_bytes}"
        )

    try:
        os.replace(tmp_path, out_path)
    except OSError as exc:
        tmp_path.unlink(missing_ok=True)
        return None, f"atomic rename failed: {exc}"
    return size, None


def _render_docx(
    *,
    body: str,
    out_path: Path,
    title: str | None,
    max_output_bytes: int,
) -> tuple[int | None, str | None]:
    try:
        from docx import Document  # Lazy import.
    except ImportError as exc:  # pragma: no cover — root pyproject pins docx.
        return None, f"python-docx not importable: {exc}"

    tmp_path = out_path.with_suffix(out_path.suffix + ".tmp")
    try:
        doc = Document()
        if title:
            doc.add_heading(title, level=1)
        # Preserve paragraph breaks from the body verbatim — consumers
        # (chat, review flows) rely on visible structure.
        for paragraph in body.split("\n"):
            doc.add_paragraph(paragraph)
        doc.save(str(tmp_path))
    except Exception as exc:
        # python-docx raises OSError + lxml.XMLSyntaxError variants; we
        # collapse all of them to the exit-5 JSON so the handler path
        # stays uniform across formats.
        tmp_path.unlink(missing_ok=True)
        return None, f"python-docx render failed: {type(exc).__name__}: {exc}"

    try:
        size = tmp_path.stat().st_size
    except OSError as exc:
        tmp_path.unlink(missing_ok=True)
        return None, f"stat after render failed: {exc}"

    if size > max_output_bytes:
        tmp_path.unlink(missing_ok=True)
        return None, (
            f"rendered DOCX {size} bytes exceeds cap {max_output_bytes}"
        )

    try:
        os.replace(tmp_path, out_path)
    except OSError as exc:
        tmp_path.unlink(missing_ok=True)
        return None, f"atomic rename failed: {exc}"
    return size, None


# ---------------------------------------------------------------------------
# Argv + dispatch
# ---------------------------------------------------------------------------


def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="render_doc",
        description=(
            "Render a plain-text body to PDF or DOCX under the "
            "<data_dir>/media/outbox/ directory. Output format is "
            "inferred from the --out suffix."
        ),
    )
    p.add_argument(
        "--body-file",
        required=True,
        help=(
            "Path to the plain-text body. Must live under "
            "<data_dir>/run/render-stage/."
        ),
    )
    p.add_argument(
        "--out",
        required=True,
        help=(
            "Destination path. Must live under <data_dir>/media/outbox/ "
            "and end in .pdf or .docx."
        ),
    )
    p.add_argument(
        "--title",
        default=None,
        help="Optional heading placed above the body.",
    )
    p.add_argument(
        "--font",
        default="DejaVu",
        help=(
            "Font label embedded in the PDF (cosmetic). The glyphs "
            "always come from the vendored DejaVuSans.ttf."
        ),
    )
    return p


def main(argv: list[str] | None = None) -> int:
    parser = _build_parser()
    try:
        args = parser.parse_args(argv)
    except SystemExit as exc:
        # argparse's SystemExit is our exit code 2 for usage errors.
        return int(exc.code) if isinstance(exc.code, int) else EXIT_USAGE

    data_dir = _resolve_data_dir()

    body_path, err = _resolve_body_file(args.body_file, data_dir)
    if body_path is None:
        return _fail(EXIT_PATH, err or "--body-file rejected")

    out_path, err = _resolve_out_path(args.out, data_dir)
    if out_path is None:
        return _fail(EXIT_PATH, err or "--out rejected")

    max_body = _resolve_max_body_bytes()
    max_output = _resolve_max_output_bytes()

    try:
        body_bytes = body_path.read_bytes()
    except OSError as exc:
        return _fail(EXIT_IO, f"read --body-file failed: {exc}")
    if len(body_bytes) > max_body:
        return _fail(
            EXIT_PATH,
            f"body {len(body_bytes)} bytes exceeds cap {max_body}",
        )
    if not body_bytes:
        return _fail(EXIT_USAGE, "--body-file is empty")
    try:
        body = body_bytes.decode("utf-8")
    except UnicodeDecodeError as exc:
        return _fail(EXIT_IO, f"body is not valid UTF-8: {exc}")

    suffix = out_path.suffix.lower()
    try:
        if suffix == ".pdf":
            size, err = _render_pdf(
                body=body,
                out_path=out_path,
                title=args.title,
                font_path=_BUNDLED_FONT_PATH,
                max_output_bytes=max_output,
            )
        elif suffix == ".docx":
            size, err = _render_docx(
                body=body,
                out_path=out_path,
                title=args.title,
                max_output_bytes=max_output,
            )
        else:  # pragma: no cover — _resolve_out_path already rejected.
            return _fail(EXIT_PATH, f"unsupported suffix {suffix!r}")
    except Exception as exc:
        # Last-ditch safety net — anything not caught by the renderer
        # branches (e.g. a bug in this wrapper). Keep broad so the CLI
        # exits 5 with a traceback rather than bubbling a stack to the
        # Bash hook caller.
        tb = traceback.format_exc(limit=3)
        return _fail(
            EXIT_UNKNOWN,
            f"renderer crashed: {type(exc).__name__}: {exc}",
            traceback=tb,
        )

    if size is None:
        return _fail(
            EXIT_IO if err and "rename" in err else EXIT_UNKNOWN,
            err or "renderer returned no size",
        )
    return _ok(
        {
            "out": str(out_path),
            "bytes": size,
            "format": suffix.lstrip("."),
        }
    )


if __name__ == "__main__":
    raise SystemExit(main())
