"""extract_doc CLI — local text extraction from PDF / DOCX / XLSX / RTF / TXT.

Invoked via Bash allowlist as `python tools/extract_doc/main.py <path>
[--max-chars N] [--pages N-M]`. Dispatch by suffix (case-insensitive).

Security model (phase-7 §2.9):

- `defusedxml.defuse_stdlib()` is called at module load time so every
  `xml.etree.ElementTree` / `xml.sax` / `xml.dom.minidom` call inside
  python-docx or openpyxl is routed through defusedxml's hardened
  parsers. This blocks billion-laughs / quadratic-blowup / external-
  entity (XXE) attacks that hostile DOCX / XLSX authors can embed.
- DOCX and XLSX are ZIPs: before handing them to python-docx / openpyxl
  we inspect the zip directory and refuse to extract if the sum of
  declared uncompressed sizes exceeds `_ZIP_UNCOMPRESSED_CAP` (64 MB —
  3x the 20 MB input cap, sufficient for legit docs, tight enough to
  block zip-bomb amplification).
- Input file size is capped at 20 MB by default
  (`MEDIA_EXTRACT_MAX_INPUT_BYTES` env override) before any parse.
- Path is resolved + `is_file()` checked; symlinks are followed
  intentionally so the caller can pass an inbox file through the
  adapter chain.

Output: single JSON line on stdout; errors go to stderr as
`{"ok": false, "error": "..."}`.

Exit codes:
  0  ok
  2  usage (argparse)
  3  validation (path / size / format / zip-bomb / encrypted)
  4  I/O (open / read / zip-corrupt)
  5  unknown (unexpected exception — caller should log and retry)
"""

from __future__ import annotations

import argparse
import json
import os
import sys
import zipfile
from pathlib import Path
from typing import Any, Final

# Install defusedxml patches into xml.* stdlib modules BEFORE python-docx /
# openpyxl import — both libraries cache the patched parser factories at
# import time. We also import defusedxml modules ourselves for the one
# place we parse XML directly (none at the moment, but keeping the import
# makes the security posture explicit + survives future refactors).
from defusedxml import defuse_stdlib  # type: ignore[import-untyped]
from defusedxml.common import (  # type: ignore[import-untyped]  # noqa: F401
    DefusedXmlException,  # re-exported for tests
)

defuse_stdlib()

# The rest of the imports follow; they are deliberately below the
# `defuse_stdlib()` call even though ruff would normally group them.
import pypdf  # noqa: E402
from docx import Document as _DocxDocument  # noqa: E402
from openpyxl import load_workbook  # type: ignore[import-untyped]  # noqa: E402
from striprtf.striprtf import rtf_to_text  # type: ignore[import-untyped]  # noqa: E402

# --- constants --------------------------------------------------------------

EXIT_OK: Final = 0
EXIT_USAGE: Final = 2
EXIT_VALIDATION: Final = 3
EXIT_IO: Final = 4
EXIT_UNKNOWN: Final = 5

_DEFAULT_MAX_INPUT_BYTES: Final = 20_000_000  # mirrors MediaSettings.extract_max_input_bytes
_ENV_MAX_INPUT_BYTES: Final = "MEDIA_EXTRACT_MAX_INPUT_BYTES"

# 3x input cap — legit .docx / .xlsx rarely exceed this after
# decompression; zip-bomb authors need orders-of-magnitude amplification
# to be interesting, so 64 MB is a comfortable ceiling.
_ZIP_UNCOMPRESSED_CAP: Final = 64 * 1024 * 1024

_MAX_CHARS_HARD_CAP: Final = 2_000_000  # 2 MB text out — protects dispatch_reply + history
_DEFAULT_MAX_CHARS: Final = 200_000  # plenty for a model turn; caller can request more

_SUPPORTED_SUFFIXES: Final = frozenset({".pdf", ".docx", ".xlsx", ".rtf", ".txt"})

# --- XML-entity attack scan (phase-7 fix-pack D2) ---------------------------
#
# python-docx (OPC package) + openpyxl parse XML parts via lxml's
# ``XMLParser(resolve_entities=False)`` — NOT via stdlib ``xml.*``.
# ``defusedxml.defuse_stdlib()`` cannot see those code paths, so a
# hostile DOCX/XLSX with a ``<!DOCTYPE>`` / ``<!ENTITY>`` declaration
# reached the parser unnoticed. lxml's ``resolve_entities=False`` keeps
# the attack non-leaky (entities stay unresolved, no external fetch),
# but the CLI returned exit 0 with empty text — identical to a legit
# empty document — which is a silent-acceptance gap.
#
# Fix: before handing a ZIP-backed Office file to python-docx /
# openpyxl, scan every XML entry in the zip for a DOCTYPE or ENTITY
# declaration. If present, reject with ``EXIT_VALIDATION`` and an
# explicit structured-log event. Keeps defense-in-depth (the
# ``_guard_zip_bomb`` check still fires first for archive-level
# amplification), adds explicit XXE / billion-laughs rejection that
# surfaces in the CLI's exit code.
#
# Bytes-level scan rather than parse: (a) parsing the hostile XML
# with ``defusedxml.lxml`` would require an extra dependency and
# defeat the point of scanning before any parser touches the
# payload; (b) a simple bytes-level substring scan is O(file_size),
# bounded by ``_ZIP_UNCOMPRESSED_CAP`` via the zip-bomb guard; (c)
# case-insensitive match covers XML spec variants ``<!DOCTYPE`` and
# ``<!doctype`` alike, while the marker characters themselves
# cannot legitimately appear in any XML part content (they're only
# valid at the prolog level).
_XML_ENTITY_MARKERS: Final[tuple[bytes, ...]] = (b"<!DOCTYPE", b"<!ENTITY")

# Cap the declared-size of a single XML part we're willing to scan —
# an individual XML part >32 MB is unheard of in legitimate Office
# files and a hostile crafted part could otherwise push RAM. 32 MB is
# half of ``_ZIP_UNCOMPRESSED_CAP`` so the sum across all XML parts
# is still under the zip-bomb ceiling.
_XML_PART_SCAN_CAP: Final = 32 * 1024 * 1024

# XML part suffixes found inside DOCX / XLSX packages. We scan parts
# matching these — the payload surface an attacker could manipulate.
# Other entries in an OPC zip (e.g. ``word/media/image1.png``,
# ``[Content_Types].xml``) are either non-XML (images) or already
# covered by the .xml/.rels scan.
_XML_PART_SUFFIXES: Final[tuple[str, ...]] = (".xml", ".rels")


# --- output helpers ---------------------------------------------------------


def _ok(payload: dict[str, Any]) -> int:
    sys.stdout.write(json.dumps({"ok": True, **payload}, ensure_ascii=False))
    sys.stdout.write("\n")
    return EXIT_OK


def _fail(code: int, error: str, **extra: Any) -> int:
    body: dict[str, Any] = {"ok": False, "error": error}
    body.update(extra)
    sys.stderr.write(json.dumps(body, ensure_ascii=False))
    sys.stderr.write("\n")
    return code


# --- text post-processing ---------------------------------------------------


def _truncate(text: str, limit: int) -> tuple[str, bool]:
    """Tail-trim a string to `limit` characters.

    Returns (truncated_text, was_truncated). We clamp to
    `_MAX_CHARS_HARD_CAP` irrespective of what the caller asked for so a
    model can't accidentally blow the history / dispatch pipeline up.
    """
    effective = min(limit, _MAX_CHARS_HARD_CAP)
    if len(text) <= effective:
        return text, False
    # Don't split a surrogate pair; Python str indexing is by code point
    # so `text[:effective]` is always valid — no special handling needed.
    return text[:effective], True


# --- page-range parsing -----------------------------------------------------


def _parse_page_range(spec: str) -> tuple[int, int]:
    """Parse `--pages N-M` (1-based, inclusive).

    Accepts a single page (`5` → (5, 5)) or a hyphenated range
    (`2-7` → (2, 7)). Raises ValueError on malformed input.
    """
    s = spec.strip()
    if not s:
        raise ValueError("empty page spec")
    if "-" in s:
        left, _, right = s.partition("-")
        a = int(left.strip())
        b = int(right.strip())
    else:
        a = b = int(s)
    if a < 1 or b < 1:
        raise ValueError(f"pages must be >=1, got {spec!r}")
    if b < a:
        raise ValueError(f"end page {b} precedes start page {a}")
    return a, b


# --- validation -------------------------------------------------------------


def _validate_path(raw: str) -> tuple[Path, str]:
    """Resolve + verify an input path. Returns (resolved_path, suffix_lower)."""
    try:
        path = Path(raw).resolve(strict=True)
    except (OSError, RuntimeError) as exc:
        raise _RejectError(EXIT_VALIDATION, f"path resolve failed: {exc}") from exc
    if not path.is_file():
        raise _RejectError(EXIT_VALIDATION, f"not a regular file: {path}")
    suffix = path.suffix.lower()
    if suffix not in _SUPPORTED_SUFFIXES:
        raise _RejectError(
            EXIT_VALIDATION,
            f"unsupported suffix {suffix!r}; expected one of "
            + ", ".join(sorted(_SUPPORTED_SUFFIXES)),
        )
    return path, suffix


def _validate_size(path: Path) -> int:
    """Enforce the input-size cap. Returns the file size in bytes."""
    try:
        size = path.stat().st_size
    except OSError as exc:
        raise _RejectError(EXIT_IO, f"stat failed: {exc}") from exc
    cap_raw = os.environ.get(_ENV_MAX_INPUT_BYTES)
    try:
        cap = int(cap_raw) if cap_raw is not None else _DEFAULT_MAX_INPUT_BYTES
    except ValueError:
        cap = _DEFAULT_MAX_INPUT_BYTES
    if cap > 0 and size > cap:
        raise _RejectError(
            EXIT_VALIDATION,
            f"input size {size} exceeds cap {cap}",
            size=size,
            cap=cap,
        )
    return size


class _RejectError(Exception):
    """Internal control-flow carrier — converts to `_fail(code, ...)`.

    Using an exception keeps every extractor free of error-return
    plumbing; `main` converts these into stderr + exit-code at the
    outermost frame.
    """

    def __init__(self, code: int, error: str, **extra: Any) -> None:
        super().__init__(error)
        self.code = code
        self.error = error
        self.extra = extra


# --- zip-bomb guard ---------------------------------------------------------


def _guard_zip_bomb(path: Path) -> None:
    """Pre-parse guard for ZIP-backed Office formats (DOCX / XLSX).

    We open the archive index only (`ZipFile` without extracting), sum
    the declared uncompressed sizes, and refuse to proceed if the total
    exceeds `_ZIP_UNCOMPRESSED_CAP`. This blocks both naive zip bombs
    (`42.zip`-style) and practical amplification attacks that stay
    below our input cap yet explode on open.

    defusedxml guards against XML-entity expansion once the parser
    reaches individual XML parts; zip-bomb is a separate attack
    surface defusedxml cannot see, hence this explicit check.

    Fix-pack D2: we also scan every XML part inside the zip for
    ``<!DOCTYPE`` / ``<!ENTITY`` markers and reject the whole file if
    any part declares one. python-docx + openpyxl route their XML
    parsing through lxml (not stdlib), which ``defuse_stdlib()``
    cannot intercept; scanning bytes-level BEFORE any parser opens
    the file closes the silent-acceptance gap and yields a real exit
    code for observability.
    """
    try:
        with zipfile.ZipFile(path) as zf:
            total = 0
            for info in zf.infolist():
                # `file_size` is the declared uncompressed size; a
                # crafted archive may lie, but the legitimate parser
                # will refuse to extract past the declared value.
                total += max(info.file_size, 0)
                if total > _ZIP_UNCOMPRESSED_CAP:
                    raise _RejectError(
                        EXIT_VALIDATION,
                        "zip-bomb guard tripped: declared uncompressed "
                        f"size exceeds cap {_ZIP_UNCOMPRESSED_CAP}",
                        declared=total,
                        cap=_ZIP_UNCOMPRESSED_CAP,
                    )
            # D2: entity-declaration scan across every XML-like part.
            # Run AFTER the size tally so the cheap check (aggregate
            # size) rejects zip-bombs before we start reading
            # individual parts — which themselves could be huge.
            _reject_xml_entity_declarations(zf, path)
    except zipfile.BadZipFile as exc:
        raise _RejectError(EXIT_IO, f"not a valid zip archive: {exc}") from exc
    except OSError as exc:
        raise _RejectError(EXIT_IO, f"zip open failed: {exc}") from exc


def _reject_xml_entity_declarations(zf: zipfile.ZipFile, path: Path) -> None:
    """Raise ``_RejectError`` if any XML part declares ``<!DOCTYPE`` / ``<!ENTITY``.

    Scans each ``.xml`` / ``.rels`` entry in the archive up to
    ``_XML_PART_SCAN_CAP`` bytes (bounded above by the zip-bomb cap
    total already enforced). A single matching byte sequence is
    sufficient to reject the whole file — legitimate Office XML parts
    never declare a DOCTYPE or ENTITY (the OPC spec puts schema
    declarations in dedicated parts the parser reaches separately).

    Using a bytes-level substring scan rather than an XML parse is
    intentional:

    * No parser receives the hostile input; the scan happens BEFORE
      python-docx / openpyxl touch the file.
    * The markers ``<!DOCTYPE`` / ``<!ENTITY`` are uppercase by XML
      spec (case-sensitivity varies by parser, but the legitimate
      declarations are always uppercase); lowercase variants would
      be rejected by a strict XML parser anyway. We still fold the
      scan to case-insensitive (``.upper()``) to catch the `lol`
      corner-case where the attacker lowercases the marker to evade
      a naive substring check, then relies on lax-mode parsing.

    Raises:
      ``_RejectError`` with exit code ``EXIT_VALIDATION`` and the
      offending part name in the extra payload so operators can tell
      which entry tripped the guard.
    """
    for info in zf.infolist():
        name = info.filename
        lower_name = name.lower()
        # Scan `.xml`, `.rels`, and `.xml.rels`. `endswith` on tuple
        # is stdlib-idiomatic and covers both suffix families.
        if not lower_name.endswith(_XML_PART_SUFFIXES):
            continue
        declared = max(info.file_size, 0)
        if declared > _XML_PART_SCAN_CAP:
            # An XML part exceeding 32 MB is itself deeply suspicious
            # — a legitimate Office part virtually never reaches this
            # size. Treat as a rejection rather than silently skipping
            # the scan.
            raise _RejectError(
                EXIT_VALIDATION,
                f"xml part {name!r} declared size {declared} exceeds scan cap "
                f"{_XML_PART_SCAN_CAP}",
                part=name,
                declared=declared,
                cap=_XML_PART_SCAN_CAP,
            )
        try:
            # `open().read()` within the ZipFile context honours the
            # underlying `ZipExtFile` streaming — for a 32 MB cap the
            # worst-case RAM cost is 32 MB, which is well within the
            # ambient process budget. We read at most the declared
            # size (guarded above) so an attacker cannot push us past
            # the cap by lying about the compressed size.
            with zf.open(info, "r") as member:
                data = member.read(_XML_PART_SCAN_CAP + 1)
        except (OSError, zipfile.BadZipFile) as exc:
            raise _RejectError(
                EXIT_IO,
                f"xml part {name!r} read failed: {exc}",
                part=name,
            ) from exc
        if len(data) > _XML_PART_SCAN_CAP:
            # The archive lied about the declared size and expanded
            # past our scan cap. Treat identically to a declared-size
            # overflow — reject rather than trying to reason about
            # the rest of the part.
            raise _RejectError(
                EXIT_VALIDATION,
                f"xml part {name!r} actual size exceeds scan cap "
                f"{_XML_PART_SCAN_CAP}",
                part=name,
                cap=_XML_PART_SCAN_CAP,
            )
        # Case-insensitive bytes scan. ``upper()`` on bytes is
        # ASCII-only which is exactly what we want — XML declaration
        # tokens are ASCII.
        upper = data.upper()
        for marker in _XML_ENTITY_MARKERS:
            if marker in upper:
                raise _RejectError(
                    EXIT_VALIDATION,
                    (
                        f"xml part {name!r} declares {marker.decode()!r}; "
                        "DOCTYPE / ENTITY declarations are rejected to "
                        "block XXE + billion-laughs attacks (extract_doc "
                        "processes content-only XML)"
                    ),
                    part=name,
                    marker=marker.decode(),
                    source=str(path),
                )


# --- extractors -------------------------------------------------------------


def _extract_pdf(path: Path, page_range: tuple[int, int] | None) -> tuple[str, int]:
    """Extract text from a PDF. Returns (text, pages_read)."""
    try:
        reader = pypdf.PdfReader(str(path))
    except pypdf.errors.PdfReadError as exc:
        raise _RejectError(EXIT_IO, f"pdf read failed: {exc}") from exc
    except OSError as exc:
        raise _RejectError(EXIT_IO, f"pdf open failed: {exc}") from exc
    if reader.is_encrypted:
        # pypdf will happily silently return empty pages for encrypted
        # PDFs if we don't decrypt; better to fail fast than hand back
        # a deceptive empty transcript.
        raise _RejectError(
            EXIT_VALIDATION,
            "pdf is encrypted; decrypt or provide an unlocked copy",
        )
    total_pages = len(reader.pages)
    if page_range is None:
        start, end = 1, total_pages
    else:
        start, end = page_range
        if start > total_pages:
            raise _RejectError(
                EXIT_VALIDATION,
                f"page range {start}-{end} exceeds document length {total_pages}",
            )
        end = min(end, total_pages)
    parts: list[str] = []
    for i in range(start - 1, end):  # pypdf is 0-indexed
        try:
            parts.append(reader.pages[i].extract_text() or "")
        except Exception as exc:  # pypdf raises several types (pdf-parse is opaque)
            # Per-page failure is not fatal; record an inline marker.
            parts.append(f"[page {i + 1} extraction failed: {type(exc).__name__}]")
    return "\n\n".join(parts).strip(), end - start + 1


def _extract_docx(path: Path) -> tuple[str, int]:
    """Extract paragraph + table text from a DOCX. Returns (text, paragraph_count)."""
    _guard_zip_bomb(path)
    try:
        doc = _DocxDocument(str(path))
    except (KeyError, ValueError, OSError, zipfile.BadZipFile) as exc:
        # python-docx surfaces a handful of exception types for malformed
        # archives; treat all as IO so the caller can decide to retry.
        raise _RejectError(EXIT_IO, f"docx open failed: {exc}") from exc
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    # Tables — each row on its own line, cells separated by tab.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            joined = "\t".join(c for c in cells if c)
            if joined:
                parts.append(joined)
    return "\n".join(parts).strip(), len(parts)


def _extract_xlsx(path: Path) -> tuple[str, int]:
    """Extract all non-empty cells, one row per line. Returns (text, row_count)."""
    _guard_zip_bomb(path)
    try:
        wb = load_workbook(str(path), read_only=True, data_only=True)
    except (KeyError, ValueError, OSError, zipfile.BadZipFile) as exc:
        raise _RejectError(EXIT_IO, f"xlsx open failed: {exc}") from exc
    try:
        rows_out: list[str] = []
        row_count = 0
        for sheet in wb.worksheets:
            rows_out.append(f"# sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                # openpyxl yields None for blanks; stringify the rest.
                rendered = [str(c) if c is not None else "" for c in row]
                # Drop fully-empty rows to keep the transcript small.
                if any(r.strip() for r in rendered):
                    rows_out.append("\t".join(rendered).rstrip())
                    row_count += 1
        return "\n".join(rows_out).strip(), row_count
    finally:
        wb.close()


def _extract_rtf(path: Path) -> tuple[str, int]:
    """Extract plain text from an RTF file. Returns (text, 1)."""
    try:
        raw = path.read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        raise _RejectError(EXIT_IO, f"rtf read failed: {exc}") from exc
    # striprtf is best-effort; it doesn't raise on malformed input, it
    # just returns whatever it can parse out. Empty result -> likely a
    # binary RTF or heavily corrupt; we still return success so the
    # caller can decide.
    return rtf_to_text(raw, errors="ignore").strip(), 1


def _extract_txt(path: Path) -> tuple[str, int]:
    """Read a plain-text file. Returns (text, line_count)."""
    # Try UTF-8 first (dominant in practice); fall back to latin-1 which
    # never fails so we always return a usable string.
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw = path.read_text(encoding="latin-1")
    except OSError as exc:
        raise _RejectError(EXIT_IO, f"txt read failed: {exc}") from exc
    return raw.strip(), raw.count("\n") + 1


# --- orchestrator -----------------------------------------------------------


def _extract(
    path: Path,
    suffix: str,
    page_range: tuple[int, int] | None,
) -> tuple[str, str, int]:
    """Dispatch by suffix. Returns (text, format_label, page_or_unit_count)."""
    if suffix == ".pdf":
        text, pages = _extract_pdf(path, page_range)
        return text, "pdf", pages
    if page_range is not None and suffix != ".pdf":
        # Non-PDF formats don't have a native page concept; silently
        # ignoring would hide a typo. Explicit rejection forces the
        # caller to remove the flag.
        raise _RejectError(
            EXIT_VALIDATION,
            f"--pages is only supported for .pdf inputs, got {suffix}",
        )
    if suffix == ".docx":
        text, units = _extract_docx(path)
        return text, "docx", units
    if suffix == ".xlsx":
        text, units = _extract_xlsx(path)
        return text, "xlsx", units
    if suffix == ".rtf":
        text, units = _extract_rtf(path)
        return text, "rtf", units
    if suffix == ".txt":
        text, units = _extract_txt(path)
        return text, "txt", units
    # _validate_path already screens the suffix; unreachable.
    raise _RejectError(EXIT_UNKNOWN, f"unreachable: suffix {suffix!r}")


# --- argparse wiring --------------------------------------------------------


def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="extract_doc",
        description=(
            "Extract plain text from PDF / DOCX / XLSX / RTF / TXT. "
            "Output is a single JSON line on stdout; errors go to stderr."
        ),
    )
    p.add_argument("path", help="Path to the input document")
    p.add_argument(
        "--max-chars",
        type=int,
        default=_DEFAULT_MAX_CHARS,
        help=(
            f"Truncate extracted text to this many characters "
            f"(default {_DEFAULT_MAX_CHARS}; hard cap {_MAX_CHARS_HARD_CAP})"
        ),
    )
    p.add_argument(
        "--pages",
        default=None,
        help="PDF-only: 1-based inclusive page range, e.g. '3-7' or '5'",
    )
    return p


def main(argv: list[str] | None = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)
    if args.max_chars <= 0:
        return _fail(EXIT_USAGE, "--max-chars must be positive")
    try:
        page_range = _parse_page_range(args.pages) if args.pages is not None else None
    except ValueError as exc:
        return _fail(EXIT_USAGE, f"--pages parse: {exc}")

    try:
        path, suffix = _validate_path(args.path)
        size = _validate_size(path)
        text, fmt, units = _extract(path, suffix, page_range)
    except _RejectError as rej:
        return _fail(rej.code, rej.error, **rej.extra)
    except Exception as exc:  # last-resort guard — stdlib/lib code paths vary
        return _fail(
            EXIT_UNKNOWN,
            f"unexpected error: {type(exc).__name__}: {exc}",
        )

    text, truncated = _truncate(text, args.max_chars)
    return _ok(
        {
            "path": str(path),
            "format": fmt,
            "units": units,
            "size_bytes": size,
            "chars": len(text),
            "truncated": truncated,
            "text": text,
        }
    )


if __name__ == "__main__":
    sys.exit(main())
